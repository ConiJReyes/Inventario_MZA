# pylint: disable=no-member, missing-function-docstring, trailing-whitespace, line-too-long, ungrouped-imports, redefined-outer-name, unused-argument, missing-module-docstring, missing-final-newline
from django.contrib.auth import logout
from django.db.models import Sum
from docx import Document
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth import authenticate, login
from django.contrib.auth.models import User, Group
from django.views.decorators.http import require_POST
from .models import Movimiento, Notificacion, Producto, Kit, KitComponente, Lote
from .models import Proveedor
from django.contrib import messages
from django.core.mail import send_mail
from django.contrib import messages
# Umbral que define cuántas unidades mínimas se consideran "stock bajo"
STOCK_UMBRAL_BAJO = 50
from django.db.models import Q

def home(request):
    return render(request, 'inventario_app/home.html')

def login_view(request):
    if request.method == 'POST':
        correo = request.POST['correo']
        contrasena = request.POST['contrasena']      
        user = authenticate(request, username=correo, password=contrasena)
        if user is not None:
            login(request, user)  
            return redirect('dashboard')  
        else:
            return render(request, 'inventario_app/login.html', {'error': 'Credenciales incorrectas'}) 

    return render(request, 'inventario_app/login.html')  

def logout_view(request):
    logout(request)  
    return redirect('login')  

@staff_member_required
def registro_view(request):
    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        correo = request.POST.get('correo')
        contrasena = request.POST.get('contrasena')
        confirmar = request.POST.get('confirmar')
        rol = request.POST.get('rol')  

        # Verificar que las contraseñas coinciden
        if contrasena != confirmar:
            return render(request, 'inventario_app/registro.html', {'error': 'Las contraseñas no coinciden'})

        # Verificar si el correo ya está registrado
        if User.objects.filter(username=correo).exists():
            return render(request, 'inventario_app/registro.html', {'error': 'Correo ya registrado'})

        usuario = User.objects.create_user(username=nombre, email=correo, password=contrasena, first_name=nombre)

        # Asignar el grupo según el rol seleccionado
        if rol:
            grupo = Group.objects.filter(name=rol).first()
            if grupo:
                usuario.groups.add(grupo)

        return redirect('dashboard')

    # Pasar los grupos al template para el select
    grupos = Group.objects.all() 
    return render(request, 'inventario_app/registro.html', {'grupos': grupos})

def dashboard(request):
    producto_count = Producto.objects.filter(estado=True).aggregate(Sum('stock'))['stock__sum'] or 0
    kit_count = Kit.objects.count()
    usuarios = User.objects.all()
    lotes_count = Lote.objects.count()
    usuarios_count = usuarios.count()
    return render(request, 'inventario_app/dashboard.html', {'producto_count': producto_count, 'usuarios_count':usuarios_count, 'kit_count': kit_count, 'lotes_count': lotes_count})

#PARA EL CRUD DE LOS PRODUCTOS
#VISTA PARA LISTAR LOS PRODUCTOS
@login_required
def productos_inventario(request):
    productos = Producto.objects.filter(estado=True)
    kits = Kit.objects.all().order_by('nombre').prefetch_related('kitcomponentes__producto')
    lotes = Lote.objects.all()
    return render(request, 'inventario_app/productos_inventario.html', {'productos': productos, 'kits': kits, 'lotes':lotes})

#VISTA PARA EDITARLOS
@login_required
def editar_producto_separado(request, pk):
    producto = get_object_or_404(Producto, pk=pk)

    if request.method == 'POST':
        # Asignar valores a los campos del producto
        producto.nombre = request.POST.get('nombre', producto.nombre)
        producto.descripcion = request.POST.get('descripcion', producto.descripcion)
        producto.stock = int(request.POST.get('stock', producto.stock))
        producto.etiqueta = request.POST.get('etiqueta', producto.etiqueta)

        # Procesar la fecha de vencimiento, asegurándose de que el formato sea correcto
        fv = request.POST.get('fecha_vencimiento', None)
        if fv:
            producto.fecha_vencimiento = fv


        if 'imagen' in request.FILES:
            producto.imagen = request.FILES['imagen']
    
        producto.save()

        return redirect('productos_inventario')

    return render(request, 'inventario_app/editar_inventario.html', {'producto': producto})

#VISTA PARA ELIMINAR U PRODUCTO
@login_required
@require_POST
def eliminar_producto(request, pk):
    producto = get_object_or_404(Producto, pk=pk)
    producto.estado = False  # marcar como inactivo
    producto.save()          # guardar cambios
    return redirect('productos_inventario')


#PAGINA DEL CONFIGURACION
@login_required
def configuracion(request):
    # Obtener los grupos del usuario logueado
    grupos_usuario = request.user.groups.values_list('name', flat=True)

    if request.method == 'POST':
        user_id = request.POST.get('user_id')
        usuario = get_object_or_404(User, pk=user_id)
        rol_seleccionado = request.POST.getlist('groups')  # Obtener los roles seleccionados

        # Lógica para prevenir que el administrador cambie su rol
        if 'Administrador' in rol_seleccionado and 'Administrador' not in [grupo.name for grupo in usuario.groups.all()]:
            return render(request, 'inventario_app/configuracion.html', {
                'error': 'No puedes asignar el rol de Administrador.',
                'usuarios': User.objects.all(),
                'grupos': Group.objects.all(),
                'grupos_usuario': grupos_usuario  # Pasamos los grupos al contexto
            })
        
        # Verificar si los datos realmente cambiaron
        # Comparar los datos actuales con los datos del formulario
        cambios = False
        if usuario.first_name != request.POST.get('first_name'):
            usuario.first_name = request.POST.get('first_name')
            cambios = True
        if usuario.email != request.POST.get('email'):
            usuario.email = request.POST.get('email')
            cambios = True

        # Verificar si los grupos fueron modificados
        roles_actuales = [grupo.id for grupo in usuario.groups.all()]
        if sorted(rol_seleccionado) != sorted(roles_actuales):
            usuario.groups.set(Group.objects.filter(id__in=rol_seleccionado))
            cambios = True

        if cambios:
            usuario.save()
            mensaje = "Cambios guardados exitosamente."
        else:
            mensaje = "No hay cambios que guardar."

        return render(request, 'inventario_app/configuracion.html', {
            'mensaje': mensaje,
            'usuarios': User.objects.all(),
            'grupos': Group.objects.all(),
            'grupos_usuario': grupos_usuario  # Pasamos los grupos al contexto
        })

    # Si el método es GET, solo mostramos el formulario
    return render(request, 'inventario_app/configuracion.html', {
        'usuarios': User.objects.all(),
        'grupos': Group.objects.all(),
        'grupos_usuario': grupos_usuario,  # Para los usuarios logueados
    })


@login_required
@require_POST
def eliminar_usuario( pk):
    usuario = get_object_or_404(User, pk=pk)
    usuario.delete()
    return redirect('configuracion')


# MOVIMIENTOS
@login_required
def movimientos(request):
    producto_agotado = None
    toast_alerta = None
    productos_stock_bajo = None

    # ----- Filtros GET -----
    fecha_inicio = request.GET.get('fecha_inicio', '')
    fecha_fin = request.GET.get('fecha_fin', '')
    buscar = request.GET.get('producto_buscar', '').strip()
    tm_filter = request.GET.get('tipo_movimiento_filter', '')

    if request.method == 'POST':
        # --- LÓGICA EXISTENTE DE CREACIÓN DE MOVIMIENTO ---
        tipo_elemento   = request.POST.get('tipo_elemento')
        producto_id     = request.POST.get('producto')
        tipo_movimiento = request.POST.get('tipo_movimiento')
        cantidad        = int(request.POST.get('cantidad'))
        motivo_movimiento = request.POST.get('motivo_movimiento')

        producto = get_object_or_404(Producto, id=producto_id)

        # Alerta: salida > 50
        if tipo_movimiento == 'salida' and cantidad > 50:
            mensaje = (
                f"Se ha realizado un movimiento de salida de más de 50 "
                f"unidades del producto {producto.nombre}."
            )
            Notificacion.objects.create(usuario=request.user, mensaje=mensaje)
            toast_alerta = (
                f"¡Alerta! Se ha realizado una salida de más de 50 unidades "
                f"de {producto.nombre}."
            )

        # Alerta: stock bajo tras movimiento
        nuevo_stock = producto.stock - cantidad if tipo_movimiento=='salida' else producto.stock + cantidad
        if nuevo_stock < STOCK_UMBRAL_BAJO:
            mensaje_stock = (
                f"El stock de {producto.nombre} es bajo. Quedan {nuevo_stock} unidades."
            )
            Notificacion.objects.create(usuario=request.user, mensaje=mensaje_stock)
            productos_stock_bajo = Producto.objects.filter(
                stock__lt=STOCK_UMBRAL_BAJO, estado=True
            )

        # Ajustar stock de la pieza
        if tipo_movimiento == 'entrada':
            producto.stock += cantidad
        else:  # 'salida'
            if producto.stock >= cantidad:
                producto.stock -= cantidad
                if producto.stock == 0:
                    producto.estado = False
                    producto_agotado = producto.nombre
            else:
                return render(request, 'inventario_app/movimientos.html', {
                    'error': 'No hay suficiente stock para la salida',
                    'productos': Producto.objects.filter(estado=True),
                    'movimientos': Movimiento.objects.all().order_by('-fecha'),
                    'movimientos_count': Movimiento.objects.count(),
                    'productos_stock_bajo': None,
                    'toast_alerta': None,
                    'fecha_inicio': fecha_inicio,
                    'fecha_fin': fecha_fin,
                    'producto_buscar': buscar,
                    'tipo_movimiento_filter': tm_filter,
                })
        producto.save()

        Movimiento.objects.create(
            producto=producto,
            tipo_movimiento=tipo_movimiento,
            cantidad=cantidad,
            motivo_movimiento=motivo_movimiento,
            usuario=request.user
        )

        # tras POST, recargamos listado **sin filtros GET**
        return render(request, 'inventario_app/movimientos.html', {
            'movimientos': Movimiento.objects.all().order_by('-fecha'),
            'movimientos_count': Movimiento.objects.count(),
            'productos': Producto.objects.filter(estado=True),
            'productos_stock_bajo': productos_stock_bajo,
            'toast_alerta': toast_alerta,
            'fecha_inicio': '',
            'fecha_fin': '',
            'producto_buscar': '',
            'tipo_movimiento_filter': '',
        })

    # --- SI ES GET: aplicamos los filtros en la consulta ---
    qs = Movimiento.objects.all().order_by('-fecha')

    if fecha_inicio:
        qs = qs.filter(fecha__date__gte=fecha_inicio)
    if fecha_fin:
        qs = qs.filter(fecha__date__lte=fecha_fin)
    if buscar:
        qs = qs.filter(
            Q(producto__nombre__icontains=buscar) |
            Q(lote__numero_lote__icontains=buscar)
        )
    if tm_filter in ('entrada', 'salida'):
        qs = qs.filter(tipo_movimiento=tm_filter)

    return render(request, 'inventario_app/movimientos.html', {
        'movimientos': qs,
        'movimientos_count': qs.count(),
        'productos': Producto.objects.filter(estado=True),
        'productos_stock_bajo': None,
        'toast_alerta': None,
        'fecha_inicio': fecha_inicio,
        'fecha_fin': fecha_fin,
        'producto_buscar': buscar,
        'tipo_movimiento_filter': tm_filter,
        'motivo_movimiento' : Movimiento.motivo_movimiento
    })

#HTML stock bajo
@login_required
def notificaciones_view(request):
    notificaciones = Notificacion.objects.filter(usuario=request.user).order_by('-fecha')
    return render(request, 'inventario_app/notificaciones.html', {'notificaciones': notificaciones})

def eliminar_notificacion(request, noti_id):
    noti = get_object_or_404(Notificacion, id=noti_id, usuario=request.user)
    noti.delete()
    return redirect('notificaciones')

#REPORTES 
@login_required
def reportes_view(request):
    movimientos = Movimiento.objects.all().order_by('-fecha')

    return render(request, 'inventario_app/reportes.html', {
        'movimientos': movimientos,
        'movimientos_count': movimientos.count(),
    })


# Generar el reporte en Word
@login_required
def generar_reporte_word(request):
    movimientos = Movimiento.objects.all()
    doc = Document()
    doc.add_heading('Reporte de Movimientos de Inventario', 0)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    # Encabezados de la tabla
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Producto'
    hdr_cells[1].text = 'Tipo de Movimiento'
    hdr_cells[2].text = 'Cantidad'
    hdr_cells[3].text = 'Fecha'

    # Llenar la tabla con los movimientos
    for movimiento in movimientos:
        row_cells = table.add_row().cells
        row_cells[0].text = movimiento.producto.nombre
        row_cells[1].text = movimiento.tipo_movimiento
        row_cells[2].text = str(movimiento.cantidad)
        row_cells[3].text = str(movimiento.fecha)

    # Crear la respuesta HTTP con el archivo Word
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="movimientos.docx"'

    # Guardar el archivo Word en la respuesta HTTP
    doc.save(response)
    return response


#REGISTRAR PRODUCTOS
@login_required
def registrar_producto(request):
    piezas = Producto.objects.filter(estado=True)

    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        descripcion = request.POST.get('descripcion')
        stock = request.POST.get('stock')
        imagen = request.FILES.get('imagen')  
        etiqueta = request.POST.get('etiquetas') 

        # Crear el nuevo producto
        Producto.objects.create(
            nombre=nombre,
            descripcion=descripcion,
            stock=stock,
            imagen=imagen,
            etiqueta=etiqueta
        )

        return redirect('productos_inventario')  
    return render(request, 'inventario_app/registrar_producto.html',{ 'piezas':piezas })  

#PROVEEDORES
@login_required
def proveedores_view(request):
    if request.method == 'POST':
        # Obtener los datos del formulario
        nombre = request.POST.get('nombre')
        contacto = request.POST.get('contacto')
        telefono = request.POST.get('telefono')
        email = request.POST.get('email')

        # Crear un nuevo proveedor
        Proveedor.objects.create(
            nombre=nombre,
            contacto=contacto,
            telefono=telefono,
            email=email
        )
        return redirect('proveedores')  
    # Obtener todos los proveedores registrados
    proveedores = Proveedor.objects.all()
    return render(request, 'inventario_app/proveedores.html', {'proveedores': proveedores})

# Editar proveedor existente
@login_required
def editar_proveedor(request, pk):
    proveedor = get_object_or_404(Proveedor, pk=pk)

    if request.method == 'POST':
        proveedor.nombre = request.POST.get('nombre')
        proveedor.contacto = request.POST.get('contacto')
        proveedor.telefono = request.POST.get('telefono')
        proveedor.email = request.POST.get('email')
        proveedor.save()

        return redirect('proveedores')  # Redirigir a la lista de proveedores después de la edición

    return render(request, 'inventario_app/editar_proveedor.html', {'proveedor': proveedor})

# Eliminar proveedor
@login_required
def eliminar_proveedor(request, pk):
    proveedor = get_object_or_404(Proveedor, pk=pk)
    proveedor.delete()  # Eliminar proveedor de la base de datos
    return redirect('proveedores')  # Redirigir a la lista de proveedores después de eliminar

def crear_kit(request):
    if request.method == 'POST':
        nombre = request.POST.get('kit_nombre')
        descripcion = request.POST.get('kit_descripcion')
        pieza_ids = request.POST.getlist('pieza_ids[]')
        cantidades = request.POST.getlist('cantidades[]')

        # Validaciones
        if not nombre:
            messages.error(request, "El nombre del kit es obligatorio.")
            return redirect('inventario_app/registrar_producto.html')

        if not pieza_ids or not cantidades or len(pieza_ids) != len(cantidades):
            messages.error(request, "Debe seleccionar piezas y cantidades válidas.")
            return redirect('inventario_app/registrar_producto.html')

        try:
            # Crear el kit
            kit, created = Kit.objects.get_or_create(nombre=nombre, descripcion=descripcion)

            # Crear los componentes del kit (relación con las piezas)
            for pieza_id, cantidad in zip(pieza_ids, cantidades):
                pieza = Producto.objects.get(id=pieza_id)
                cantidad = int(cantidad)
                if cantidad <= 0:
                    continue
                KitComponente.objects.create(kit=kit, producto=pieza, cantidad=cantidad)

            # Calcular el stock del kit basado en las piezas
            cantidades_posibles = []
            for componente in kit.kitcomponentes.all():
                # Calculamos cuántos kits podemos hacer con el stock de cada pieza
                cantidad_posible = componente.producto.stock // componente.cantidad
                cantidades_posibles.append(cantidad_posible)

            # El stock del kit es el mínimo de los valores posibles según las piezas
            kit.stock = min(cantidades_posibles) if cantidades_posibles else 0
            kit.save()

            # Actualizar el stock basado en las piezas
            kit.actualizar_stock()

            messages.success(request, f"Se ha creado el kit '{kit.nombre}' con un stock de {kit.stock} kits.")
            return redirect('dashboard/')

        except Producto.DoesNotExist:
            messages.error(request, "Alguna de las piezas seleccionadas no existe.")
            return redirect('inventario_app/registrar_producto.html')

        except Exception as e:
            messages.error(request, f"Error al crear el kit: {str(e)}")
            return redirect('inventario_app/registrar_producto.html')

    else:
        piezas = Producto.objects.filter(estado=True).order_by('nombre')
        return render(request, 'inventario_app/registrar_producto.html', {'piezas': piezas})

def registrar_lote(request):
    if request.method == 'POST':
        producto_id = request.POST.get('producto')
        numero_lote = request.POST.get('numero_lote')
        cantidad = int(request.POST.get('cantidad'))
        fecha_vencimiento = request.POST.get('fecha_vencimiento')
        marca = request.POST.get('marca')

        # Validaciones
        if not producto_id or not numero_lote or cantidad <= 0:
            messages.error(request, "Todos los campos son obligatorios y la cantidad debe ser mayor que cero.")
            return redirect('registrar_lote')  # Redirige de nuevo al formulario

        try:
            producto = Producto.objects.get(id=producto_id)
            lote = Lote.objects.create(
                producto=producto,
                numero_lote=numero_lote,
                cantidad=cantidad,
                fecha_vencimiento=fecha_vencimiento,
                marca=marca
            )
            lote.producto.actualizar_stock()  # Actualizar el stock del producto

            messages.success(request, f"Lote '{numero_lote}' registrado exitosamente.")
            return redirect('productos_inventario')  # Redirige a la vista de inventarios

        except Producto.DoesNotExist:
            messages.error(request, "Producto no encontrado.")
            return redirect('registrar_lote')

    else:
        productos = Producto.objects.all()  # Obtener todos los productos
        return render(request, 'inventario_app/registrar_producto.html', {'piezas': productos})
    
def detalles_movimiento(request, id_movimiento):
    # Obtener el movimiento específico
    movimiento = get_object_or_404(Movimiento, id=id_movimiento)
    # Pasar el movimiento a la plantilla
    return render(request, 'inventario_app/detalles_movimiento.html', {'movimiento': movimiento})